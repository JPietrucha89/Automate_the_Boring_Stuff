import docx
import os
from pathlib import Path

# creating blank Word file
d = docx.Document()

# adding new paragraphs
d.add_paragraph('Hello! This is first paragraph.')
d.add_paragraph('This is another paragraph.')

# adding new runs to existing paragraph
p = d.paragraphs[0]
p.add_run('This is new run in first paragraph.')
p.runs[1].bold = True

# NOTE Unfortunately there is no way to insert paragraphs or runs between other paragraphs or runs. In fact both .add() methods are like .append()
d.save(Path(os.getcwd(), 'working_with_Word', 'created_file.docx'))

# ADDING HEADERS
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# ADDING PAGE BREAK
# To add a line break (rather than starting a whole new paragraph), you can call the add_break() method on the Run object you want to have the break appear after. If you want to add a page break instead, you need to pass the value docx.enum.text.WD_BREAK.PAGE as a lone argument to add_break()
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

# ADDING IMAGES
# The optional width and height keyword arguments will set the width and height of the image in the document. If left out, the width and height will default to the normal size of the image.
doc = docx.Document()
doc.add_picture(Path(os.getcwd(), 'working_with_Word', 'zophie.png'), width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))
doc.save('word_with_image.docx')
