import docx
import os
from pathlib import Path

d = docx.Document(Path(os.getcwd(), 'working_with_Word',
                  'demo.docx'))  # create Document object
print(d.paragraphs)  # print paragraphs from Document object
print(d.paragraphs[0].text)  # print text contents of first paragraph

p = d.paragraphs[1]
print(p.runs)  # every paragraph contains one or more runs. New run starts whenever there is change in style of text (for example bolding, italicizing)
print('* Printing all runs from second paragraph...')
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print('* Checking BOLD attribute in all runs from second paragraph...')
print(p.runs[0].bold)
print(p.runs[1].bold)
print(p.runs[2].bold)
print(p.runs[3].bold)

print('* Checking ITALIC attribute in all runs from second paragraph...')
print(p.runs[0].italic)
print(p.runs[1].italic)
print(p.runs[2].italic)
print(p.runs[3].italic)

# CHANGING FORMATTING
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

# CHANGING STYLING
print('* Printing styles of each paragraph ...')
for paragraph in d.paragraphs:
    print(paragraph.style)
# Let's change style of second paragraph
d.paragraphs[1].style = 'Title'

# SAVING WORD FILE
d.save(Path(os.getcwd(), 'working_with_Word', 'demo_altered.docx'))
